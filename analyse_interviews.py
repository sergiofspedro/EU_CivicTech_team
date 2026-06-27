"""
analyse_interviews.py
Phase 1 analysis of interview transcripts using the Claude API.

For each .docx transcript, applies the Phase 1 analytical prompt and saves
the structured output as a .docx file in the OUTPUT_FOLDER.

Setup:
  pip install anthropic python-docx
  set ANTHROPIC_API_KEY=your_key_here   (Windows PowerShell)

Usage:
  # Test on one file:
  python analyse_interviews.py --file "IFP20_Thijs_Berman_EPD_ENG.docx"

  # Process all .docx in INTERVIEWS_FOLDER:
  python analyse_interviews.py --all
"""

import os
import re
import sys
import argparse
import zipfile
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── CONFIGURATION ─────────────────────────────────────────────────────────────
INTERVIEWS_FOLDER = "interviews"          # folder with .docx transcripts
OUTPUT_FOLDER     = "interview_analyses"  # where output .docx files are saved
MODEL             = "claude-opus-4-8"     # best model for deep qualitative analysis
MAX_TOKENS        = 8000
# ───────────────────────────────────────────────────────────────────────────────

BOOK_CONTEXT = """
BOOK TITLE (working title):
"Contribution of Networks to the Intensification of Democratic Practices in Europe.
Main Pathways for Changing Dynamics"

EXTENDED SYNTHESIS:
This volume is produced within the final months of the Nets4Dem project
("Innovating democracy, together", https://nets4dem.eu) and in dialogue with ScalDem
(https://scaledem.eu). Nets4Dem was born from an initial consortium of 11 organisations
spread across Europe and the United States, with the aim of creating a network of networks
— a hub for collaboration and advanced research in democratic innovation, civic deliberation,
participatory democracy, and citizenship education. Today the network covers 200 cities,
50 think tanks and universities, 38 European countries, and 320 civil society organisations.

The book is conceived as a "bridging tool" for a transition moment, providing instruments
to the original consortium and new members to critically reflect on themselves and their
interactions, acquiring a reflexive approach to their future networking. It builds on:
- Two Social Network Analyses conducted among consortium members (2024 and 2025)
- A Knowledge Database on Democratic Innovation drawing on 400+ EU-funded research projects
- ~40 qualitative interviews with leaders of network organisations across Europe

The central themes of the book include:
- Types of networks and network of networks in civil society
- Democratic intensification vs. democratic backsliding and erosion
- Network resilience, formation, governance, and transformation
- Internal governance, membership, legal forms, and funding
- Transnational and multi-level civil society networks in Europe
- The tension between collaboration and competition among civil society organisations
- The role of networks in shaping European democratic governance

The interviews were conducted with leaders of civil society networks, platforms, alliances,
federations, and umbrella organisations working on participatory democracy, democratic
innovation, and civic education in Europe and globally.
"""

PHASE1_PROMPT = """
Instruction for the model:
I will provide you with the transcript of one interview. The content of this interview is
referred to a network organisation that works with the goal of deepening democratic
practices and fighting democratic backsliding in Europe. Your task is to analyse the text
in a very accurate, comprehensive, and faithful way, and to extract all the relevant
themes/topics that emerge from the interviewee's speech.

OBJECTIVE:
I want a thorough reconstruction of the interview content, with particular attention to:
- main themes
- sub-themes
- recurring concepts
- tensions, dilemmas, or contradictions
- concrete examples, cited cases, names of actors, projects, institutions, or territories
- organisational, political, relational, methodological, financial and normative dimensions

WHAT YOU MUST PRODUCE:
For each theme identified, you must provide:

1. THEME TITLE — A short but precise label.
2. THEME SUMMARY — A brief summary of 3–6 lines.
3. DETAILED EXPLANATION — A more extended explanation of what the interviewee says about
   that theme, preserving as much as possible the original meaning of the discourse.
4. SPECIFIC ELEMENTS MENTIONED — examples; actors involved; places or contexts;
   organisational practices; problems or opportunities; any contradictions or nuances.
5. KEY QUOTATIONS OR FORMULATIONS — If useful, include short and significant formulations
   from the interviewee, without over-paraphrasing and without inventing anything.
   Put them in "…"

METHODOLOGICAL RULES:
- Do not invent information that is not present in the text.
- Do not oversimplify: preserve the complexity of the discourse.
- Do not lose nuances, hesitations, ambivalences, or contradictions.
- If a theme appears several times in different parts of the interview, merge the references
  and explain how it develops across the discourse.
- If some passages are unclear or ambiguous, state that clearly.
- Maintain an analytical, faithful, and orderly tone.
- Avoid personal judgments or interpretations not supported by the text.
- If a theme is very important but secondary with respect to the main topic, include it anyway.
- If the discourse touches on implicit aspects that can be reconstructed with good evidence,
  you may label them as an [IMPLICIT THEME].

IMPORTANT CONSTRAINTS:
- The response must be complete but readable.
- Do not rely only on overly general categories.
- If the transcript is long, prioritise depth and full coverage rather than excessive summarisation.
- Begin your response with a brief INTERVIEW OVERVIEW (3–5 lines): who is the interviewee,
  what organisation, main role, and the general character/tone of the interview.
- Then list all themes in numbered order.
- At the end, include a SHORT THEMATIC MAP (bullet list) summarising all identified themes
  and their relationships.
"""


# ── Helpers ────────────────────────────────────────────────────────────────────

def extract_text_from_docx(path: str) -> str:
    """Extract plain text from a .docx file."""
    try:
        with zipfile.ZipFile(path) as z:
            with z.open("word/document.xml") as f:
                xml = f.read().decode("utf-8")
        text = re.sub(r"<[^>]+>", " ", xml)
        text = re.sub(r"\s+", " ", text).strip()
        text = text.replace("&apos;", "'").replace("&amp;", "&") \
                   .replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"')
        return text
    except Exception as e:
        print(f"  [ERROR] Could not read {path}: {e}")
        return ""


def analyse_with_claude(interview_text: str, filename: str) -> str:
    """Send interview to Claude for Phase 1 analysis."""
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    system = (
        "You are an expert qualitative researcher specialising in civil society, "
        "democratic innovation, and network governance in Europe. You produce rigorous, "
        "structured, and faithful analyses of interview transcripts."
    )

    user_message = f"""
{BOOK_CONTEXT}

---

{PHASE1_PROMPT}

---

INTERVIEW TRANSCRIPT (file: {filename}):

{interview_text}

---

Please now produce the full Phase 1 analysis as instructed above.
"""

    print(f"  Sending to Claude API ({MODEL}) ...", flush=True)
    message = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": user_message}],
    )
    return message.content[0].text


def save_analysis_docx(analysis_text: str, filename: str, interview_text: str):
    """Save the analysis as a formatted .docx file."""
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    stem = os.path.splitext(filename)[0]
    out_path = os.path.join(OUTPUT_FOLDER, f"ANALYSIS_{stem}.docx")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading(f"Phase 1 Analysis", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_heading(f"Interview: {filename}", level=2)
    sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    # Parse and render the analysis
    # Split on lines and detect structure
    lines = analysis_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Detect theme headers (e.g. "THEME 1:", "## THEME", numbered themes)
        if re.match(r'^#+\s+', line):
            level = len(re.match(r'^(#+)', line).group(1))
            text = re.sub(r'^#+\s+', '', line)
            h = doc.add_heading(text, level=min(level + 1, 4))
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif re.match(r'^THEME\s+\d+', line, re.IGNORECASE) or \
             re.match(r'^\d+\.\s+THEME', line, re.IGNORECASE) or \
             re.match(r'^\*\*THEME', line, re.IGNORECASE):
            clean = re.sub(r'\*+', '', line).strip()
            h = doc.add_heading(clean, level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif re.match(r'^(THEME TITLE|THEME SUMMARY|DETAILED EXPLANATION|'
                      r'SPECIFIC ELEMENTS|KEY QUOTATIONS|INTERVIEW OVERVIEW|'
                      r'SHORT THEMATIC MAP|THEMATIC MAP|\[IMPLICIT THEME\])', line, re.IGNORECASE) or \
             re.match(r'^\*\*(THEME TITLE|THEME SUMMARY|DETAILED EXPLANATION|'
                      r'SPECIFIC ELEMENTS|KEY QUOTATIONS)', line, re.IGNORECASE):
            clean = re.sub(r'\*+', '', line).strip().rstrip(':')
            h = doc.add_heading(clean, level=3)
            h.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r'^\s*[-•]\s*', '', line)
            # Handle **bold** within bullet
            _add_formatted_run(p, text)
        elif re.match(r'^\d+\.\s+', line) and not re.match(r'^\d+\.\s+THEME', line, re.IGNORECASE):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r'^\d+\.\s+', '', line)
            _add_formatted_run(p, text)
        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, line)

        i += 1

    doc.save(out_path)
    print(f"  Saved → {os.path.abspath(out_path)}")
    return out_path


def _add_formatted_run(paragraph, text: str):
    """Add text to paragraph, handling **bold** markdown."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def process_file(path: str):
    filename = os.path.basename(path)
    print(f"\nProcessing: {filename}")

    # Check if already done
    stem = os.path.splitext(filename)[0]
    out_path = os.path.join(OUTPUT_FOLDER, f"ANALYSIS_{stem}.docx")
    if os.path.exists(out_path):
        print(f"  [SKIP] Analysis already exists: {out_path}")
        return

    interview_text = extract_text_from_docx(path)
    if not interview_text:
        print(f"  [ERROR] Empty transcript — skipping.")
        return
    print(f"  Transcript: {len(interview_text.split())} words")

    analysis = analyse_with_claude(interview_text, filename)
    print(f"  Analysis: {len(analysis.split())} words")

    save_analysis_docx(analysis, filename, interview_text)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Phase 1 interview analysis via Claude API")
    parser.add_argument("--file", help="Path to a single .docx transcript to analyse")
    parser.add_argument("--all", action="store_true",
                        help=f"Process all .docx files in {INTERVIEWS_FOLDER}/")
    args = parser.parse_args()

    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("[ERROR] ANTHROPIC_API_KEY environment variable not set.")
        print("  Windows: $env:ANTHROPIC_API_KEY = 'sk-ant-...'")
        sys.exit(1)

    if args.file:
        if not os.path.exists(args.file):
            print(f"[ERROR] File not found: {args.file}")
            sys.exit(1)
        process_file(args.file)

    elif args.all:
        if not os.path.isdir(INTERVIEWS_FOLDER):
            print(f"[ERROR] Folder '{INTERVIEWS_FOLDER}' not found.")
            print(f"  Create it and place your .docx transcripts inside.")
            sys.exit(1)
        files = sorted(
            os.path.join(INTERVIEWS_FOLDER, f)
            for f in os.listdir(INTERVIEWS_FOLDER)
            if f.lower().endswith(".docx")
        )
        if not files:
            print(f"[ERROR] No .docx files found in '{INTERVIEWS_FOLDER}/'")
            sys.exit(1)
        print(f"Found {len(files)} transcripts to process.\n")
        for path in files:
            process_file(path)

    else:
        parser.print_help()
        print("\nExamples:")
        print('  python analyse_interviews.py --file "IFP20_Thijs_Berman_EPD_ENG.docx"')
        print('  python analyse_interviews.py --all')


if __name__ == "__main__":
    main()
